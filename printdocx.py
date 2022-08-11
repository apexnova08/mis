import os
import sqlite3
import docx
from docx.shared import Pt
from openpyxl import load_workbook
from datetime import date

import misc

dbPath = directory = os.getcwd() + "\mis\mis.db"

def PrintBgyCertificate(pk, purpose):
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/brgycertificate/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(14)
    for paragraph in doc.paragraphs:

        # Officials Left
        
        if (paragraph.text == "CHAIRMAN"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD A"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad Appropriation'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD VAWC"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad VAWC'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD EP"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad EP'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD PO"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad PO'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD CL"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad CL'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD ECT"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad ECAT'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD HR"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad HR'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "SK CHAIRMAN"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'SK Chairman'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "SECRETARY"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Secretary'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "TREASURER"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Treasurer'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
            
        # Content

        if ("RESIDENT NAME" in paragraph.text):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run("This is to certify that ")
            cursor.execute("SELECT * FROM residents WHERE resPK = " + pk)
            for x in cursor:
                run = paragraph.add_run(misc.FMSNameConvert(x[1], x[2], x[3], x[4]))
                run.bold = True
                run = paragraph.add_run(" who is presently residing at ")
                run = paragraph.add_run(x[11] + " " + x[12] + " Street,")
                run.bold = True
            run = paragraph.add_run(" within the jurisdiction of Barangay 40 - Gumamela, Cavite City.")

        if ("Issued on: DATE" in paragraph.text):
            paragraph.style = doc.styles['Normal']
            paragraph.text = "Issued on: "
            run = paragraph.add_run(misc.GetMonth(str(date.today()).split("-")[1]) + " " + str(date.today()).split("-")[2] + ", " + str(date.today()).split("-")[0])
            run.bold = True

        if (paragraph.text == "Purpose"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            if (purpose != ""):
                run = paragraph.add_run("Purpose: ")
                run = paragraph.add_run(purpose)
                run.bold = True
    
    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run = paragraph.add_run(x[2])
                    run.bold = True
                    run.underline = True
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "Barangay-Certificate-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/brgycertificate/' + filename)
    os.system('start mis/docs/brgycertificate/' + filename)



def PrintBgyClearance(pk, purpose):
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/brgyclearance/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    for paragraph in doc.paragraphs:

        # Officials Left

        if (paragraph.text == "CHAIRMAN"):
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
            for x in cursor:
                run2 = paragraph.add_run(x[2])
                run2.bold = True
                run2.underline = True
                font2 = run2.font
                font2.name = 'Calibri'
                font2.size = Pt(14)
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD A"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad Appropriation'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD VAWC"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad VAWC'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD EP"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad EP'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD PO"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad PO'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD CL"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad CL'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD ECT"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad ECAT'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD HR"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad HR'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "SK CHAIRMAN"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'SK Chairman'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "SECRETARY"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Secretary'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "TREASURER"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Treasurer'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)

        # Content

        if ("RESIDENT NAME" in paragraph.text):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run("This is to certify that ")
            cursor.execute("SELECT * FROM residents WHERE resPK = " + pk)
            for x in cursor:
                run = paragraph.add_run(misc.FMSNameConvert(x[1], x[2], x[3], x[4]))
                run.bold = True
                run = paragraph.add_run(" who is presently residing at ")
                run = paragraph.add_run(x[11] + " " + x[12] + " Street,")
                run.bold = True
                if (x[20] == "Yes"):
                    run = paragraph.add_run(" is a registered and accredited member of this Barangay.")
                else:
                    run = paragraph.add_run(" is an accredited member of this Barangay.")

        if ("DATE." in paragraph.text):
            paragraph.style = doc.styles['Normal']
            paragraph.text = "This certification is issued upon the request of subject this "
            run = paragraph.add_run(misc.GetMonth(str(date.today()).split("-")[1]) + " " + str(date.today()).split("-")[2] + ", " + str(date.today()).split("-")[0])
            run.bold = True

        if (paragraph.text == "Purpose"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            if (purpose != ""):
                run = paragraph.add_run("Purpose: ")
                run = paragraph.add_run(purpose)
                run.bold = True
    
    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run2 = paragraph.add_run(x[2])
                    run2.bold = True
                    run2.underline = True
                    font2 = run2.font
                    font2.name = 'Calibri'
                    font2.size = Pt(14)
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "Barangay-Clearance-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/brgyclearance/' + filename)
    os.system('start mis/docs/brgyclearance/' + filename)

def PrintBsnClearance(rName, rAddress, bName, bAddress):
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/brgybusinessclearance/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)
    for paragraph in doc.paragraphs:

        # Content

        if (paragraph.text == "BUSINESS NAME"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run(bName.upper())
            run.bold = True
            run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "BUSINESS ADDRESS"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run(bAddress.upper())
            run.bold = True
            run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "OWNER NAME"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run(rName.upper())
            run.bold = True
            run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "OWNER ADDRESS"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run(rAddress.upper())
            run.bold = True
            run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)

        # Date

        if ("DAY day of MONTH, YEAR" in paragraph.text):
            paragraph.style = doc.styles['Normal']
            paragraph.text = "Issued this "
            run = paragraph.add_run(str(date.today()).split("-")[2])
            run.bold = True
            run = paragraph.add_run(misc.GetOrdinal(str(date.today()).split("-")[2]))
            run.bold = True
            run.font.subscript = True
            run = paragraph.add_run(" day of ")
            run = paragraph.add_run(misc.GetMonth(str(date.today()).split("-")[1]) + " " + str(date.today()).split("-")[0])
            run.bold = True
    
    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run = paragraph.add_run(x[2])
                    run.bold = True
                    run.underline = True
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "Business-Clearance-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/brgybusinessclearance/' + filename)
    os.system('start mis/docs/brgybusinessclearance/' + filename)

def PrintBgyIndigency(pk, purpose):
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/brgyindigency/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    for paragraph in doc.paragraphs:

        # Officials Left

        if (paragraph.text == "CHAIRMAN"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
            for x in cursor:
                run2 = paragraph.add_run(x[2])
                run2.bold = True
                run2.underline = True
                font2 = run2.font
                font2.name = 'Calibri'
                font2.size = Pt(14)
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD A"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad Appropriation'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD VAWC"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad VAWC'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD EP"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad EP'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD PO"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad PO'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD CL"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad CL'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD ECT"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad ECAT'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "KAGAWAD HR"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Kagawad HR'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "SK CHAIRMAN"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'SK Chairman'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "SECRETARY"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Secretary'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)
        if (paragraph.text == "TREASURER"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            cursor.execute("SELECT * FROM officials WHERE oPosition = 'Treasurer'")
            for x in cursor:
                run = paragraph.add_run(x[2])
                run.bold = True
                run.underline = True
            paragraph.paragraph_format.space_after = Pt(0)

        # Content

        if ("RESIDENT NAME" in paragraph.text):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run("This is to certify that ")
            cursor.execute("SELECT * FROM residents WHERE resPK = " + pk)
            for x in cursor:
                run = paragraph.add_run(misc.FMSNameConvert(x[1], x[2], x[3], x[4]))
                run.bold = True
                run = paragraph.add_run(" who is presently residing at ")
                run = paragraph.add_run(x[11] + " " + x[12] + " Street,")
                run.bold = True
            run = paragraph.add_run(" is a registered and accredited member of this Barangay.")

        if ("Issued on: DATE" in paragraph.text):
            paragraph.style = doc.styles['Normal']
            paragraph.text = "Issued on: "
            run = paragraph.add_run(misc.GetMonth(str(date.today()).split("-")[1]) + " " + str(date.today()).split("-")[2] + ", " + str(date.today()).split("-")[0])
            run.bold = True

        if (paragraph.text == "Purpose"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            if (purpose != ""):
                run = paragraph.add_run("Purpose: ")
                run = paragraph.add_run(purpose)
                run.bold = True
    
    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run2 = paragraph.add_run(x[2])
                    run2.bold = True
                    run2.underline = True
                    font2 = run2.font
                    font2.name = 'Calibri'
                    font2.size = Pt(14)
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "Barangay-Indigency-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/brgyindigency/' + filename)
    os.system('start mis/docs/brgyindigency/' + filename)

def PrintPWDList():
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/listpwd/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    data = []
    num = 1

    cursor.execute("SELECT * FROM residents WHERE resPWD = 'Yes' ORDER BY resSName ASC, resFName ASC, resNameExt ASC, resMName ASC")
    for x in cursor:
        ne = ""
        if (x[4] != ""):
            ne = ", " + str(x[4])
        data.append((num, x[1], x[2] + ne,  misc.GetMI(x[3]), x[11] + " " + x[12] + " St., Caridad, Cavite City"))
        num = num + 1

    # Date

    for paragraph in doc.paragraphs:
        if (paragraph.text == "DATE"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run2 = paragraph.add_run(misc.GetMonth(str(date.today()).split("-")[1]) + " " + str(date.today()).split("-")[2] + ", " + str(date.today()).split("-")[0])
            run2.bold = True
            font2 = run2.font
            font2.name = 'Calibri'
            font2.size = Pt(14)

    # Content

    for no, sname, fname, mi, address in data:
    
        row = doc.tables[0].add_row().cells
        
        row[0].text = str(no)
        row[1].text = sname
        row[2].text = fname
        row[3].text = mi
        row[4].text = address

    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run2 = paragraph.add_run(x[2])
                    run2.bold = True
                    run2.underline = True
                    font2 = run2.font
                    font2.name = 'Calibri'
                    font2.size = Pt(14)
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "PWD-List-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/listpwd/' + filename)
    os.system('start mis/docs/listpwd/' + filename)

def PrintSeniorList():
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/listseniorcitizen/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    data = []
    num = 1

    cursor.execute("SELECT * FROM residents WHERE resAge >= 60 ORDER BY resSName ASC, resFName ASC, resNameExt ASC, resMName ASC")
    for x in cursor:
        ne = ""
        if (x[4] != ""):
            ne = ", " + str(x[4])
        data.append((num, x[1], x[2] + ne,  misc.GetMI(x[3]), misc.GetMonth(x[7].split("-")[1]) + " " + x[7].split("-")[2] + ", " + x[7].split("-")[0]))
        num = num + 1

    # Content

    for no, sname, fname, mi, bdate in data:
    
        row = doc.tables[0].add_row().cells
        
        row[0].text = str(no)
        row[1].text = sname
        row[2].text = fname
        row[3].text = mi
        row[4].text = bdate

    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run2 = paragraph.add_run(x[2])
                    run2.bold = True
                    run2.underline = True
                    font2 = run2.font
                    font2.name = 'Calibri'
                    font2.size = Pt(14)
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "Senior-List-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/listseniorcitizen/' + filename)
    os.system('start mis/docs/listseniorcitizen/' + filename)

def PrintPensionerList():
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/listpensioners/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    data = []
    num = 1

    cursor.execute("SELECT * FROM residents WHERE resPensioner = 'Yes' ORDER BY resSName ASC, resFName ASC, resNameExt ASC, resMName ASC")
    for x in cursor:
        ne = ""
        if (x[4] != ""):
            ne = ", " + str(x[4])
        data.append((num, x[1], x[2] + ne, misc.GetMI(x[3]), x[11] + " " + x[12] + " St., Caridad, Cavite City"))
        num = num + 1

    # Content

    for no, sname, fname, mi, address in data:
    
        row = doc.tables[0].add_row().cells
        
        row[0].text = str(no)
        row[1].text = sname
        row[2].text = fname
        row[3].text = mi
        row[4].text = address

    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run2 = paragraph.add_run(x[2])
                    run2.bold = True
                    run2.underline = True
                    font2 = run2.font
                    font2.name = 'Calibri'
                    font2.size = Pt(14)
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "Pensioners-List-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/listpensioners/' + filename)
    os.system('start mis/docs/listpensioners/' + filename)

def PrintCustomList(searchString, title):
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    doc = docx.Document("mis/docs/listcustom/template.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    data = []
    num = 1

    cursor.execute(searchString + " ORDER BY resSName ASC, resFName ASC, resNameExt ASC, resMName ASC")
    for x in cursor:
        ne = ""
        if (x[4] != ""):
            ne = ", " + str(x[4])
        data.append((num, x[1], x[2] + ne, misc.GetMI(x[3]), x[11] + " " + x[12] + " St., Caridad, Cavite City"))
        num = num + 1

    # Content

    for no, sname, fname, mi, address in data:
    
        row = doc.tables[0].add_row().cells
        
        row[0].text = str(no)
        row[1].text = sname
        row[2].text = fname
        row[3].text = mi
        row[4].text = address

    # Title

    for paragraph in doc.paragraphs:
        if (paragraph.text == "TITLE"):
            paragraph.style = doc.styles['Normal']
            paragraph.text = ""
            run = paragraph.add_run(title.upper())
            run.bold = True
            
    # Signature

    for table in doc.tables:
        for paragraph in table.rows[0].cells[0].paragraphs:
            if (paragraph.text == "BARANGAY CHAIRMAN NAME"):
                paragraph.style = doc.styles['Normal']
                paragraph.text = ""
                cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
                for x in cursor:
                    run2 = paragraph.add_run(x[2])
                    run2.bold = True
                    run2.underline = True
                    font2 = run2.font
                    font2.name = 'Calibri'
                    font2.size = Pt(14)
                paragraph.paragraph_format.space_after = Pt(0)
    db.close()
    filename = "Resident-Custom-List-" + misc.GetNowString() + ".docx"
    doc.save('mis/docs/listcustom/' + filename)
    os.system('start mis/docs/listcustom/' + filename)

def PrintHousehold(block, lot):
    db = sqlite3.connect(dbPath)
    cursor = db.cursor()
    workbook = load_workbook(filename="mis/docs/household/template.xlsx")
    
    sheet = workbook.active
    
    # Top

    sheet["C8"] = "Block " + block + " Lot " + lot
    sheet["O8"] = misc.GetMonth(str(date.today()).split("-")[1]) + " " + str(date.today()).split("-")[2] + ", " + str(date.today()).split("-")[0]

    # Content

    i = 13
    cursor.execute("SELECT * FROM residents WHERE resHouseholdLot = '" + lot + "' AND resHouseholdBlock = '" + block + "' ORDER BY resSName ASC, resFName ASC, resNameExt ASC, resMName ASC")
    for x in cursor:
        sheet["A" + str(i)] = x[1]
        sheet["C" + str(i)] = x[2]
        sheet["F" + str(i)] = x[3]
        sheet["H" + str(i)] = x[4]
        sheet["I" + str(i)] = x[11]
        sheet["J" + str(i)] = x[12]
        sheet["N" + str(i)] = x[8]
        sheet["O" + str(i)] = x[7]
        sheet["P" + str(i)] = x[6]
        sheet["Q" + str(i)] = x[9]
        sheet["R" + str(i)] = x[10]
        sheet["S" + str(i)] = x[23]
        i = i + 1

    # Bottom

    cursor.execute("SELECT * FROM residents WHERE resHouseholdLot = '" + lot + "' AND resHouseholdBlock = '" + block + "' AND resHead = 'Yes'")
    for x in cursor:
        sheet["A31"] = misc.FMSNameConvert(x[1], x[2], x[3], x[4])
    
    cursor.execute("SELECT * FROM officials WHERE oPosition = 'Secretary'")
    for x in cursor:
        sheet["I31"] = x[2]

    cursor.execute("SELECT * FROM officials WHERE oPosition = 'Chairman'")
    for x in cursor:
        sheet["O31"] = x[2]

    db.close()
    namefile = "Household-" + misc.GetNowString() + ".xlsx"
    workbook.save(filename='mis/docs/household/' + namefile)
    os.system('start mis/docs/household/' + namefile)